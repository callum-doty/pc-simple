"""
AI service - handles document analysis using LLM APIs
Consolidates OCR, text extraction, and AI analysis into a single service
Now integrated with PromptManager for sophisticated analysis

Provider split: Anthropic (Claude) performs all document analysis and OCR.
OpenAI is used solely for embeddings (text-embedding-3-small). There is no
Gemini or OpenAI-as-analysis-provider path — the pipeline commits to this
split rather than carrying unused multi-provider branching.
"""

import asyncio
import html
import io
import logging
import re
from typing import Dict, Any, List, Optional, Tuple, AsyncGenerator, Generator
import json
import base64
from pathlib import Path
import fitz  # PyMuPDF
import anthropic
import openai
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from config import get_settings
from services.storage_service import StorageService
from services.taxonomy_service import TaxonomyService
from services.prompt_manager import PromptManager

logger = logging.getLogger(__name__)
settings = get_settings()


class AIService:
    """Unified AI service for document analysis with PromptManager integration"""

    # Bump whenever the OCR system/user prompt wording changes in a way that
    # could change what gets extracted (not just cosmetic). Stored in
    # Document.file_metadata.ocr_prompt_version for every OCR'd document (pdf
    # or image) so a shift in extracted_text quality/shape can be correlated
    # with a known prompt change instead of investigated as a data mystery.
    #
    # Changelog:
    #   1 — original prompt ("preserve original formatting as much as
    #       possible" with no output-format constraint). Undocumented at the
    #       time; not recorded on any document processed under it.
    #   2 — (2026-07-14) explicit plain-text-only instruction: forbids
    #       markdown (**bold**, # headers) and HTML entities (&nbsp;), which
    #       version 1 allowed the model to produce (e.g. representing a
    #       whitespace gap in a print slug line as repeated &nbsp; entities).
    OCR_PROMPT_VERSION = 2

    def __init__(self, db: Session):
        self.db = db
        self.storage_service = StorageService()
        self.taxonomy_service = TaxonomyService(db=self.db)
        self.prompt_manager = PromptManager(taxonomy_service=self.taxonomy_service)

        # Initialize AI clients with explicit parameter handling.
        # anthropic_client performs all analysis/OCR; openai_client is embeddings-only.
        self.anthropic_client = None
        self.openai_client = None

        if settings.anthropic_api_key:
            try:
                self.anthropic_client = anthropic.Anthropic(
                    api_key=settings.anthropic_api_key
                )
            except Exception as e:
                logger.warning(f"Failed to initialize Anthropic client: {str(e)}")
                self.anthropic_client = None

        if settings.openai_api_key:
            try:
                self.openai_client = openai.OpenAI(api_key=settings.openai_api_key)
            except Exception as e:
                logger.warning(f"Failed to initialize OpenAI client: {str(e)}")
                self.openai_client = None

        # Anthropic is the sole document-analysis/OCR provider.
        self.ai_provider = self._determine_ai_provider()

    def _determine_ai_provider(self) -> str:
        """Anthropic is the sole analysis/OCR provider; OpenAI is embeddings-only."""
        if self.anthropic_client:
            return "anthropic"
        logger.warning(
            "No AI provider configured. AI analysis will be disabled. Please set ANTHROPIC_API_KEY."
        )
        return "none"

    async def analyze_document(
        self, file_path: str, filename: str, analysis_type: str = "unified"
    ) -> Dict[str, Any]:
        """
        Complete document analysis pipeline with multiple analysis options

        Args:
            file_path: Path to the document file
            filename: Name of the document
            analysis_type: Type of analysis to perform
                - "unified": Single comprehensive analysis (default)
                - "modular": Multiple specialized analyses
                - "metadata": Core metadata only
                - "classification": Classification only
                - "entities": Entity extraction only
                - "text": Text extraction only
                - "design": Design elements only
                - "keywords": Taxonomy keywords only
                - "communication": Communication focus only

        Returns:
            Consolidated analysis results
        """
        try:
            logger.info(f"Starting {analysis_type} analysis for document: {filename}")

            # Step 1: Get file content
            file_content = await self.storage_service.get_file(file_path)
            if not file_content:
                raise ValueError(f"Could not retrieve file content for {filename}")

            # Step 2: Determine file type and extract text
            file_type = self._get_file_type(filename)
            extracted_text = await self._extract_text(file_content, file_type, filename)

            # Image and PDF types send the rendered page/image itself to the model,
            # so empty extracted_text is fine there. Every other type relies entirely
            # on extracted_text — if it's empty, there's nothing to analyze, and
            # calling the model anyway produces a fabricated result grounded in
            # nothing but the filename. Fail loudly instead.
            if not extracted_text.strip() and file_type not in ("image", "pdf"):
                raise ValueError(
                    f"No text could be extracted from '{filename}' (file_type={file_type}). "
                    f"The file may be empty, corrupted, or in an unsupported format."
                )

            # Step 3: Perform AI analysis based on type
            if analysis_type == "unified":
                ai_analysis = await self._perform_unified_analysis(
                    extracted_text, file_content, file_type, filename
                )
            elif analysis_type == "modular":
                ai_analysis = await self._perform_modular_analysis(
                    extracted_text, file_content, file_type, filename
                )
            else:
                ai_analysis = await self._perform_specific_analysis(
                    analysis_type, extracted_text, file_content, file_type, filename
                )

            # Step 4: Validate keyword mappings and re-extract keywords/categories
            mappings = self._extract_mappings_from_analysis(ai_analysis)
            validated_mappings = await self._validate_keyword_mappings(mappings)

            # Update the analysis with validated mappings
            if "keyword_mappings" in ai_analysis:
                ai_analysis["keyword_mappings"] = validated_mappings
            if "taxonomy_keywords" in ai_analysis and isinstance(
                ai_analysis.get("taxonomy_keywords"), dict
            ):
                ai_analysis["taxonomy_keywords"][
                    "keyword_mappings"
                ] = validated_mappings

            keywords, categories = self._extract_keywords_from_analysis(ai_analysis)

            # Step 5: Consolidate results
            result = {
                "extracted_text": extracted_text,
                "ai_analysis": ai_analysis,
                "keywords": keywords,
                "categories": categories,
                "file_type": file_type,
                "analysis_provider": self.ai_provider,
                "analysis_type": analysis_type,
                # Only pdf/image go through OCR — text/docx extraction doesn't
                # use a model prompt at all, so there's no OCR version to record.
                "ocr_prompt_version": (
                    self.OCR_PROMPT_VERSION if file_type in ("image", "pdf") else None
                ),
            }

            logger.info(f"Completed {analysis_type} analysis for document: {filename}")
            return result

        except Exception as e:
            logger.error(f"Error analyzing document {filename}: {str(e)}")
            raise

    async def _perform_unified_analysis(
        self, extracted_text: str, file_content: bytes, file_type: str, filename: str
    ) -> Dict[str, Any]:
        """Perform unified analysis using the comprehensive prompt"""
        try:
            # If Anthropic is not available, fall back to modular analysis
            if not self.anthropic_client:
                logger.warning(
                    "Anthropic client not available for unified analysis. Falling back to modular analysis."
                )
                return await self._perform_modular_analysis(
                    extracted_text, file_content, file_type, filename
                )

            # Get the unified analysis prompt
            prompt_data = await self.prompt_manager.get_unified_analysis_prompt(
                filename
            )

            # Prepare image data if it's an image or PDF
            image_data = None
            if file_type in ["image", "pdf"]:
                image_data = self._prepare_image_data(file_content, file_type)

            # Add extracted text to the prompt
            enhanced_prompt = self._enhance_prompt_with_text(
                prompt_data["user"], extracted_text
            )

            # Call the AI service
            analysis_result = await self._call_anthropic_api_with_system(
                prompt_data["system"], enhanced_prompt, image_data
            )

            # Ensure the final output has a consistent structure
            if "document_analysis" not in analysis_result:
                analysis_result["document_analysis"] = {
                    "summary": analysis_result.get("summary", "No summary available"),
                    "document_type": analysis_result.get("document_type", "unknown"),
                    "campaign_type": analysis_result.get("campaign_type", "unknown"),
                    "document_tone": analysis_result.get("document_tone", "neutral"),
                }

            if "error" not in analysis_result:
                analysis_result["prompt_version"] = self.prompt_manager.PROMPT_VERSION
            return analysis_result

        except Exception as e:
            logger.error(f"Error in unified analysis: {str(e)}")
            return {"error": str(e)}

    async def _perform_modular_analysis(
        self, extracted_text: str, file_content: bytes, file_type: str, filename: str
    ) -> Dict[str, Any]:
        """Perform modular analysis using multiple specialized prompts"""
        try:
            results = {}

            # Prepare image data once
            image_data = None
            if file_type in ["image", "pdf"]:
                image_data = self._prepare_image_data(file_content, file_type)

            # Step 1: Core metadata
            metadata_result = await self._run_analysis_prompt(
                self.prompt_manager.get_core_metadata_prompt(filename),
                extracted_text,
                image_data,
            )
            results["metadata"] = metadata_result

            # Step 2: Classification (using metadata context)
            classification_result = await self._run_analysis_prompt(
                self.prompt_manager.get_classification_prompt(
                    filename, metadata_result
                ),
                extracted_text,
                image_data,
            )
            results["classification"] = classification_result

            # Step 3: Entity extraction
            entity_result = await self._run_analysis_prompt(
                self.prompt_manager.get_entity_prompt(filename, metadata_result),
                extracted_text,
                image_data,
            )
            results["entities"] = entity_result

            # Step 4: Text extraction
            text_result = await self._run_analysis_prompt(
                self.prompt_manager.get_text_extraction_prompt(
                    filename, metadata_result
                ),
                extracted_text,
                image_data,
            )
            results["text_extraction"] = text_result

            # Step 5: Design elements (only for visual documents)
            if file_type in ["image", "pdf"]:
                design_result = await self._run_analysis_prompt(
                    self.prompt_manager.get_design_elements_prompt(
                        filename, metadata_result
                    ),
                    extracted_text,
                    image_data,
                )
                results["design_elements"] = design_result

            # Step 6: Taxonomy keywords
            keyword_result = await self._run_analysis_prompt(
                await self.prompt_manager.get_taxonomy_keyword_prompt(
                    filename, metadata_result
                ),
                extracted_text,
                image_data,
            )
            results["taxonomy_keywords"] = keyword_result

            # Step 7: Communication focus
            communication_result = await self._run_analysis_prompt(
                self.prompt_manager.get_communication_focus_prompt(
                    filename, metadata_result
                ),
                extracted_text,
                image_data,
            )
            results["communication_focus"] = communication_result

            # Consolidate into a unified document_analysis structure
            if "metadata" in results and "classification" in results:
                # Extract summary and other details from the 'document_analysis' block within the 'metadata' result
                doc_analysis_data = results.get("metadata", {}).get(
                    "document_analysis", {}
                )
                summary_text = doc_analysis_data.get("summary", "No summary available")
                document_type = doc_analysis_data.get("document_type", "unknown")
                campaign_type = doc_analysis_data.get("campaign_type", "unknown")
                document_tone = doc_analysis_data.get("document_tone", "neutral")

                results["document_analysis"] = {
                    "summary": summary_text,
                    "document_type": document_type,
                    "campaign_type": campaign_type,
                    "document_tone": document_tone,
                }

            results["prompt_version"] = self.prompt_manager.PROMPT_VERSION
            return results

        except Exception as e:
            logger.error(f"Error in modular analysis: {str(e)}")
            return {"error": str(e)}

    async def _perform_specific_analysis(
        self,
        analysis_type: str,
        extracted_text: str,
        file_content: bytes,
        file_type: str,
        filename: str,
    ) -> Dict[str, Any]:
        """Perform a specific type of analysis"""
        try:
            # Prepare image data
            image_data = None
            if file_type in ["image", "pdf"]:
                image_data = self._prepare_image_data(file_content, file_type)

            # Get the appropriate prompt
            if analysis_type == "metadata":
                prompt_data = self.prompt_manager.get_core_metadata_prompt(filename)
            elif analysis_type == "classification":
                prompt_data = self.prompt_manager.get_classification_prompt(filename)
            elif analysis_type == "entities":
                prompt_data = self.prompt_manager.get_entity_prompt(filename)
            elif analysis_type == "text":
                prompt_data = self.prompt_manager.get_text_extraction_prompt(filename)
            elif analysis_type == "design":
                prompt_data = self.prompt_manager.get_design_elements_prompt(filename)
            elif analysis_type == "keywords":
                prompt_data = await self.prompt_manager.get_taxonomy_keyword_prompt(
                    filename
                )
            elif analysis_type == "communication":
                prompt_data = self.prompt_manager.get_communication_focus_prompt(
                    filename
                )
            else:
                raise ValueError(f"Unsupported analysis type: {analysis_type}")

            # Run the analysis
            result = await self._run_analysis_prompt(
                prompt_data, extracted_text, image_data
            )
            if "error" not in result:
                result["prompt_version"] = self.prompt_manager.PROMPT_VERSION
            return result

        except Exception as e:
            logger.error(f"Error in {analysis_type} analysis: {str(e)}")
            return {"error": str(e)}

    async def _run_analysis_prompt(
        self,
        prompt_data: Dict[str, str],
        extracted_text: str,
        image_data: Optional[str],
    ) -> Dict[str, Any]:
        """Run a single analysis prompt"""
        try:
            # Enhance prompt with extracted text
            enhanced_prompt = self._enhance_prompt_with_text(
                prompt_data["user"], extracted_text
            )

            # Call the AI service
            if self.ai_provider == "anthropic":
                return await self._call_anthropic_api_with_system(
                    prompt_data["system"], enhanced_prompt, image_data
                )
            else:
                return {"error": "No AI provider configured"}

        except Exception as e:
            logger.error(f"Error running analysis prompt: {str(e)}")
            return {"error": str(e)}

    def _enhance_prompt_with_text(self, prompt: str, extracted_text: str) -> str:
        """Enhance prompt with extracted text"""
        if extracted_text.strip():
            # Insert extracted text into the prompt
            text_section = (
                f"\n\nExtracted Text from Document:\n{extracted_text[:4000]}\n"
            )
            # Insert after the first line of the prompt
            lines = prompt.split("\n")
            if len(lines) > 1:
                lines.insert(1, text_section)
                return "\n".join(lines)
            else:
                return prompt + text_section
        return prompt

    def _get_fallback_analysis(self, filename: str, file_type: str) -> Dict[str, Any]:
        """Return a basic analysis when no AI provider is configured"""
        return {
            "document_analysis": {
                "summary": f"Document: {filename}",
                "document_type": "brochure",
                "campaign_type": "general",
                "election_year": None,
                "document_tone": "neutral",
            },
            "classification": {
                "category": "informational",
                "subcategory": None,
                "rationale": "AI analysis not available - no API keys configured",
            },
            "entities": {
                "client_name": None,
                "opponent_name": None,
                "creation_date": None,
            },
            "analysis_provider": "none",
            "file_type": file_type,
        }

    def _get_file_type(self, filename: str) -> str:
        """Determine file type from filename"""
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            return "pdf"
        elif extension in [".jpg", ".jpeg", ".png", ".tiff", ".bmp"]:
            return "image"
        elif extension in [".txt", ".md"]:
            return "text"
        elif extension in [".doc", ".docx"]:
            return "document"
        else:
            return "unknown"

    async def _extract_text(
        self, file_content: bytes, file_type: str, filename: str
    ) -> str:
        """Extract text from file based on type"""
        try:
            if file_type == "pdf":
                return await self._extract_text_from_pdf(file_content)
            elif file_type == "image":
                return await self._extract_text_from_image(file_content)
            elif file_type == "text":
                return file_content.decode("utf-8", errors="ignore")
            elif file_type == "document":
                return await self._extract_text_from_document(file_content)
            else:
                logger.warning(
                    f"Unsupported file type for text extraction: {file_type}"
                )
                return ""

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return ""

    async def _extract_text_from_pdf_generator(
        self, file_content: bytes
    ) -> AsyncGenerator[Tuple[int, str], None]:
        """
        Extract text from PDF page by page using AI-based OCR.
        Yields a tuple of (page_number, extracted_text).
        """
        doc = None
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            if doc.page_count == 0:
                logger.warning("PDF has no pages.")
                return

            for page_num in range(len(doc)):
                page = None
                pix = None
                try:
                    page = doc.load_page(page_num)
                    logger.info(f"Performing AI-based OCR on page {page_num + 1}.")
                    pix = page.get_pixmap(dpi=200)  # Lower DPI to save memory
                    img_data = pix.tobytes("png")

                    # Clean up pixmap immediately after use
                    pix = None

                    ocr_text = await self._extract_text_from_image(img_data)
                    if ocr_text.strip():
                        yield (page_num + 1, ocr_text)
                except Exception as page_error:
                    logger.error(
                        f"Error processing page {page_num + 1}: {str(page_error)}"
                    )
                    continue
                finally:
                    # Clean up page resources
                    if pix:
                        pix = None
                    if page:
                        page = None

        except Exception as e:
            logger.error(f"Error extracting text from PDF with AI OCR: {str(e)}")
        finally:
            # Always close the document
            if doc:
                doc.close()

    async def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF using the configured AI provider for OCR.
        This ensures that even image-based PDFs are processed correctly.
        """
        all_text = []
        async for page_num, ocr_text in self._extract_text_from_pdf_generator(
            file_content
        ):
            all_text.append(f"--- Page {page_num} ---\n{ocr_text}")
        return "\n\n".join(all_text)

    async def _extract_text_from_image(self, file_content: bytes) -> str:
        """Extract text from image using the configured AI provider."""
        try:
            if self.ai_provider == "none":
                logger.warning("No AI provider configured for OCR.")
                return ""

            image_data = base64.b64encode(file_content).decode("utf-8")

            system_prompt = (
                "You are an expert OCR engine. Extract all text visible in the image, "
                "accurately and completely. Output plain text only — no markdown "
                "formatting (no **bold**, no # headers, no bullet/list syntax) and no "
                "HTML entities (never write &nbsp; or similar; use a real space "
                "character). Use line breaks to reflect the reading order and layout, "
                "but represent every character as plain text exactly as printed. "
                "Only return the extracted text, with no additional comments, "
                "introductions, or summaries."
            )
            user_prompt = "Please extract all text from this image."

            raw_text = await self._call_ai_for_raw_text(
                system_prompt, user_prompt, image_data
            )
            return self._sanitize_ocr_text(raw_text)

        except Exception as e:
            logger.error(f"Error extracting text from image with AI: {str(e)}")
            return ""

    @staticmethod
    def _sanitize_ocr_text(text: str) -> str:
        """
        Defensive cleanup for OCR output. The OCR prompt asks for plain text,
        but the model doesn't always comply — e.g. representing a visual
        whitespace gap in a print-production slug line as a run of `&nbsp;`
        entities instead of a real space, or wrapping perceived headings/bold
        text in markdown syntax. This is a best-effort net on top of the
        prompt fix, not a full markdown parser.
        """
        if not text:
            return text

        # Decode HTML entities (&nbsp; -> U+00A0 non-breaking space, &amp; -> &, etc.)
        text = html.unescape(text)

        # Strip markdown bold/emphasis markers, keeping the text inside them.
        # DOTALL: OCR'd headline text often wraps a bold span across multiple lines.
        # Bold (**/__) first, then leftover single */_ italic pairs.
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"__(.+?)__", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"\*(.+?)\*", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"(?<!\w)_(.+?)_(?!\w)", r"\1", text, flags=re.DOTALL)

        # Strip ATX-style markdown headers ("# Heading" -> "Heading").
        text = re.sub(r"(?m)^#{1,6}\s+", "", text)

        # Collapse runs of horizontal whitespace (regular + non-breaking spaces)
        # down to a single space, without touching intentional line breaks.
        text = re.sub(r"[ \t ]{2,}", " ", text)

        return text.strip()

    async def _extract_text_from_document(self, file_content: bytes) -> str:
        """
        Extract text from a .docx Word document via python-docx.
        Legacy binary .doc files are not supported — python-docx raises on them,
        which surfaces as empty text and fails the document with a clear error
        rather than silently analyzing nothing.
        """
        try:
            document = DocxDocument(io.BytesIO(file_content))

            parts = [p.text for p in document.paragraphs if p.text.strip()]

            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            return "\n".join(parts)

        except Exception as e:
            logger.error(f"Error extracting text from Word document: {str(e)}")
            return ""

    def _prepare_image_data(self, file_content: bytes, file_type: str) -> Optional[str]:
        """Prepare image data for AI analysis"""
        try:
            if file_type == "image":
                # Encode image as base64
                return base64.b64encode(file_content).decode("utf-8")
            elif file_type == "pdf":
                # Convert first page of PDF to image
                doc = fitz.open(stream=file_content, filetype="pdf")
                if doc.page_count > 0:
                    page = doc[0]
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    doc.close()
                    return base64.b64encode(img_data).decode("utf-8")
            return None

        except Exception as e:
            logger.error(f"Error preparing image data: {str(e)}")
            return None

    async def _call_ai_for_raw_text(
        self, system_prompt: str, user_prompt: str, image_data: Optional[str] = None
    ) -> str:
        """Call Anthropic and return the raw text response."""
        try:
            if not self.anthropic_client:
                logger.warning("OCR not available: ANTHROPIC_API_KEY not configured.")
                return ""

            messages = [
                {
                    "role": "user",
                    "content": (
                        [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": image_data,
                                },
                            },
                            {"type": "text", "text": user_prompt},
                        ]
                        if image_data
                        else user_prompt
                    ),
                }
            ]
            response = self.anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=4000,
                system=system_prompt,
                messages=messages,
            )
            return response.content[0].text.strip()

        except Exception as e:
            logger.error(f"Error calling AI for raw text OCR: {str(e)}")
            return ""

    async def _call_anthropic_api_with_system(
        self, system_prompt: str, user_prompt: str, image_data: Optional[str] = None
    ) -> Dict[str, Any]:
        """Call Anthropic Claude API with system and user prompts"""
        try:
            messages = []

            if image_data:
                # Include image in the message
                messages.append(
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": image_data,
                                },
                            },
                            {"type": "text", "text": user_prompt},
                        ],
                    }
                )
            else:
                messages.append({"role": "user", "content": user_prompt})

            response = self.anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=3000,  # Increased for more detailed responses
                system=system_prompt,
                messages=messages,
            )

            # Parse JSON response
            response_text = response.content[0].text
            try:
                return json.loads(response_text)
            except json.JSONDecodeError:
                # If JSON parsing fails, try to extract JSON from response
                return self._extract_json_from_response(response_text)

        except Exception as e:
            logger.error(f"Error calling Anthropic API: {str(e)}")
            return {"error": str(e)}

    def _extract_json_from_response(self, response_text: str) -> Dict[str, Any]:
        """Try to extract JSON from a response that may contain additional text"""
        try:
            # Look for JSON blocks in the response
            import re

            json_pattern = r"```json\s*(.*?)\s*```"
            matches = re.findall(json_pattern, response_text, re.DOTALL)

            if matches:
                return json.loads(matches[0])

            # Try to find JSON-like content
            start_idx = response_text.find("{")
            end_idx = response_text.rfind("}")

            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                json_str = response_text[start_idx : end_idx + 1]
                return json.loads(json_str)

            # If all else fails, return raw response
            return {"raw_response": response_text}

        except Exception as e:
            logger.error(f"Error extracting JSON from response: {str(e)}")
            return {"raw_response": response_text, "extraction_error": str(e)}

    def _extract_mappings_from_analysis(
        self, ai_analysis: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Extracts keyword mappings from the analysis result."""
        if not isinstance(ai_analysis, dict):
            return []

        all_mappings = []

        # Look in the top level
        if "keyword_mappings" in ai_analysis and isinstance(
            ai_analysis["keyword_mappings"], list
        ):
            all_mappings.extend(ai_analysis["keyword_mappings"])

        # Look in the nested taxonomy_keywords block (for modular analysis)
        if "taxonomy_keywords" in ai_analysis and isinstance(
            ai_analysis["taxonomy_keywords"], dict
        ):
            nested_mappings = ai_analysis["taxonomy_keywords"].get(
                "keyword_mappings", []
            )
            if isinstance(nested_mappings, list):
                all_mappings.extend(nested_mappings)

        return all_mappings

    def _extract_keywords_from_analysis(
        self, ai_analysis: Dict[str, Any]
    ) -> Tuple[List[str], List[str]]:
        """Extract keywords and categories from AI analysis with taxonomy validation"""
        keywords = []
        categories = []
        if isinstance(ai_analysis, dict):
            # Handle keyword mappings with taxonomy validation
            if "keyword_mappings" in ai_analysis:
                mappings = ai_analysis["keyword_mappings"]
                if isinstance(mappings, list):
                    for mapping in mappings:
                        if isinstance(mapping, dict):
                            verbatim_term = mapping.get("verbatim_term")
                            canonical_term = mapping.get("mapped_canonical_term")
                            primary_category = mapping.get("mapped_primary_category")
                            if verbatim_term:
                                keywords.append(verbatim_term)
                            if canonical_term:
                                keywords.append(canonical_term)
                            if primary_category:
                                categories.append(primary_category)
            # Handle modular analysis format
            if "taxonomy_keywords" in ai_analysis:
                taxonomy_data = ai_analysis["taxonomy_keywords"]
                if (
                    isinstance(taxonomy_data, dict)
                    and "keyword_mappings" in taxonomy_data
                ):
                    mappings = taxonomy_data["keyword_mappings"]
                    if isinstance(mappings, list):
                        for mapping in mappings:
                            if isinstance(mapping, dict):
                                verbatim_term = mapping.get("verbatim_term")
                                canonical_term = mapping.get("mapped_canonical_term")
                                primary_category = mapping.get(
                                    "mapped_primary_category"
                                )

                                if verbatim_term:
                                    keywords.append(verbatim_term)
                                if canonical_term:
                                    keywords.append(canonical_term)
                                if primary_category:
                                    categories.append(primary_category)

            # Handle document_analysis format
            if "document_analysis" in ai_analysis:
                document_analysis = ai_analysis["document_analysis"]
                if isinstance(document_analysis, dict):
                    if "summary" in document_analysis and document_analysis["summary"]:
                        keywords.append(document_analysis["summary"])
                    if (
                        "document_type" in document_analysis
                        and document_analysis["document_type"]
                    ):
                        keywords.append(document_analysis["document_type"])
                    if (
                        "campaign_type" in document_analysis
                        and document_analysis["campaign_type"]
                    ):
                        keywords.append(document_analysis["campaign_type"])
                    if (
                        "document_tone" in document_analysis
                        and document_analysis["document_tone"]
                    ):
                        keywords.append(document_analysis["document_tone"])

            # Handle classification format
            if "classification" in ai_analysis:
                classification = ai_analysis["classification"]
                if isinstance(classification, dict):
                    if "category" in classification:
                        categories.append(classification["category"])
                    if (
                        "subcategory" in classification
                        and classification["subcategory"]
                    ):
                        keywords.append(classification["subcategory"])

            # Handle entities format
            if "entities" in ai_analysis:
                entities = ai_analysis["entities"]
                if isinstance(entities, dict):
                    if "client_name" in entities and entities["client_name"]:
                        keywords.append(entities["client_name"])
                    if "opponent_name" in entities and entities["opponent_name"]:
                        keywords.append(entities["opponent_name"])

            # Handle modular analysis format
            if "taxonomy_keywords" in ai_analysis:
                taxonomy_data = ai_analysis["taxonomy_keywords"]
                if (
                    isinstance(taxonomy_data, dict)
                    and "keyword_mappings" in taxonomy_data
                ):
                    mappings = taxonomy_data["keyword_mappings"]
                    if isinstance(mappings, list):
                        for mapping in mappings:
                            if isinstance(mapping, dict):
                                if (
                                    "verbatim_term" in mapping
                                    and mapping["verbatim_term"]
                                ):
                                    keywords.append(mapping["verbatim_term"])
                                if (
                                    "mapped_canonical_term" in mapping
                                    and mapping["mapped_canonical_term"]
                                ):
                                    keywords.append(mapping["mapped_canonical_term"])
                                if (
                                    "mapped_primary_category" in mapping
                                    and mapping["mapped_primary_category"]
                                ):
                                    categories.append(
                                        mapping["mapped_primary_category"]
                                    )

            # Handle communication focus
            if "communication_focus" in ai_analysis:
                comm_focus = ai_analysis["communication_focus"]
                if isinstance(comm_focus, dict):
                    if "primary_issue" in comm_focus and comm_focus["primary_issue"]:
                        keywords.append(comm_focus["primary_issue"])
                    if "messaging_strategy" in comm_focus:
                        categories.append(comm_focus["messaging_strategy"])

            # Legacy format support
            if "keywords" in ai_analysis and isinstance(ai_analysis["keywords"], list):
                keywords.extend(ai_analysis["keywords"])
            elif "key_topics" in ai_analysis and isinstance(
                ai_analysis["key_topics"], list
            ):
                keywords.extend(ai_analysis["key_topics"])

            if "categories" in ai_analysis and isinstance(
                ai_analysis["categories"], list
            ):
                categories.extend(ai_analysis["categories"])

        # Remove duplicates and None values
        keywords = list(
            set([k.strip().lower() for k in keywords if k and isinstance(k, str)])
        )
        categories = list(
            set([c.strip().lower() for c in categories if c and isinstance(c, str)])
        )

        return keywords, categories

    async def _validate_keyword_mappings(
        self, keyword_mappings: List[Dict[str, str]]
    ) -> List[Dict[str, str]]:
        """Validate and enrich keyword mappings against canonical taxonomy"""
        validated_mappings = []
        for mapping in keyword_mappings:
            canonical_term = mapping.get("mapped_canonical_term")
            if not canonical_term:
                continue
            # Get the full hierarchy for the canonical term
            hierarchy = await self.taxonomy_service.get_term_hierarchy(canonical_term)
            if hierarchy:
                enriched_mapping = {
                    "verbatim_term": mapping.get("verbatim_term"),
                    "mapped_canonical_term": hierarchy["term"],
                    "mapped_primary_category": hierarchy["primary_category"],
                    "mapped_subcategory": hierarchy["subcategory"],
                }
                validated_mappings.append(enriched_mapping)
            else:
                logger.warning(
                    f"Skipped invalid taxonomy mapping for term: {canonical_term}"
                )
        return validated_mappings

    EMBEDDING_MODEL = "text-embedding-3-small"
    EMBEDDING_VERSION = 3  # Increment when synthesis strategy changes

    _TRUSTED_CONFIDENCE = {"HIGH", "MEDIUM"}

    @staticmethod
    def build_embedding_text(
        ai_analysis: dict,
        filename: str = "",
        client_canonical: str = None,
        client_confidence: str = None,
        state: str = None,
        state_confidence: str = None,
    ) -> tuple[str, dict]:
        """
        Build a labeled, field-structured embedding string from structured AI analysis.

        Returns a (text, provenance) tuple. The provenance dict records, for each
        field that appears in the embedding string, the value used, its data source,
        and the confidence level. Store it in embedding_provenance so retrieval
        failures can be traced to specific fields without re-running the pipeline.

        Fields are ordered by semantic importance and prefixed with labels so the
        embedding model can resolve token relationships without ambiguity (e.g.
        "Florida" as state vs. topic). Handles both the full nested structure
        (holistic path) and the simplified structure from chunked processing.

        Post-processed flat columns (client_canonical, state) are gated on their
        confidence scores. A wrong value is worse than no value, so fields with
        null or unrecognised confidence are omitted rather than potentially
        embedding bad data. Controlled-vocabulary fields from ai_analysis carry
        inherently lower risk because the LLM selects from a fixed list.
        """
        # JSONB null deserializes to Python None, which is distinct from SQL NULL
        # and passes isnot(None) filters. Guard here so callers never need to check.
        if not ai_analysis:
            return filename, {}

        doc = ai_analysis.get("document_analysis", {}) or {}
        cls = ai_analysis.get("classification", {}) or {}
        ent = ai_analysis.get("entities", {}) or {}
        extracted = ai_analysis.get("extracted_text", {}) or {}
        design = ai_analysis.get("design_elements", {}) or {}
        comm = ai_analysis.get("communication_focus", {}) or {}

        summary = doc.get("summary") or ai_analysis.get("summary")

        # Canonical taxonomy terms — deduplicated, primary_issue prepended if distinct
        mappings = ai_analysis.get("keyword_mappings") or []
        canonical_terms = list({
            m["mapped_canonical_term"]
            for m in mappings
            if isinstance(m, dict) and m.get("mapped_canonical_term")
        })
        primary_issue = comm.get("primary_issue")
        if primary_issue and primary_issue not in canonical_terms:
            canonical_terms.insert(0, primary_issue)

        election_year = doc.get("election_year")

        # Gate post-processed fields on confidence — omit rather than embed bad data.
        # Controlled-vocab AI fields (document_type, category, tone, etc.) are included
        # unconditionally because the LLM picks from fixed lists.
        trusted = AIService._TRUSTED_CONFIDENCE
        client_trusted = client_canonical and (client_confidence or "").upper() in trusted
        client = client_canonical if client_trusted else ent.get("client_name")
        state_trusted = state and (state_confidence or "").upper() in trusted
        doc_state = state if state_trusted else None

        # Provenance: track every field that ends up in the embedding string.
        provenance: dict = {}

        def line(label: str, value, source: str, confidence: str = None) -> str | None:
            if not value:
                return None
            if isinstance(value, list):
                value = ", ".join(str(v) for v in value if v)
            if not value:
                return None
            provenance[label] = {"value": value, "source": source, "confidence": confidence}
            return f"{label}: {value}"

        lines = list(filter(None, [
            # Core semantic identity — highest signal, weighted first
            line("SUMMARY", summary, "llm_extraction"),
            line("MAIN_MESSAGE", extracted.get("main_message"), "llm_extraction"),
            line("CALL_TO_ACTION", extracted.get("call_to_action"), "llm_extraction"),
            # Issue vocabulary — bridges query language to taxonomy language
            line("ISSUES", ", ".join(canonical_terms) if canonical_terms else None, "taxonomy_mapping"),
            # Classification context — controlled vocabularies, low risk
            line("CATEGORY", cls.get("category"), "llm_controlled_vocab"),
            line("SUBCATEGORY", cls.get("subcategory"), "llm_controlled_vocab"),
            # Document structure — controlled vocabularies, low risk
            line("DOCUMENT_TYPE", doc.get("document_type"), "llm_controlled_vocab"),
            line("CAMPAIGN_TYPE", doc.get("campaign_type"), "llm_controlled_vocab"),
            line("TONE", doc.get("document_tone"), "llm_controlled_vocab"),
            # Audience anchor — optional prompt pass, null if not run
            line("TARGET_AUDIENCE", design.get("target_audience"), "llm_extraction"),
            # Entity anchors — confidence-gated where possible
            line(
                "CLIENT", client,
                source="canonical_pipeline" if client_trusted else "llm_extraction",
                confidence=client_confidence if client_trusted else None,
            ),
            line("OPPONENT", ent.get("opponent_name"), "llm_extraction"),
            line("CAMPAIGN", design.get("campaign_name"), "llm_extraction"),
            # Geographic and temporal context
            line(
                "STATE", doc_state,
                source="canonical_pipeline",
                confidence=state_confidence if state_trusted else None,
            ),
            line("ELECTION_YEAR", str(election_year) if election_year else None, "llm_controlled_vocab"),
        ]))

        text = "\n".join(lines) if lines else filename
        return text, provenance

    async def generate_embeddings(self, text: str) -> Optional[List[float]]:
        """Generate embeddings for text using OpenAI (embeddings-only provider)."""
        import time, asyncio
        _t = time.perf_counter()
        try:
            if not self.openai_client:
                logger.warning("Embeddings not available: OPENAI_API_KEY not configured.")
                return None
            # NOTE: openai_client is the synchronous client — run in executor to avoid
            # blocking the event loop. Replace with AsyncOpenAI to remove this overhead.
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: self.openai_client.embeddings.create(
                    model="text-embedding-3-small",
                    input=text,
                ),
            )
            logger.info(f"[PERF] openai embeddings.create: {(time.perf_counter()-_t)*1000:.0f}ms")
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Error generating embeddings with openai: {str(e)}")
            return None

    def get_ai_info(self) -> Dict[str, Any]:
        """Get information about AI configuration"""
        return {
            "ai_provider": self.ai_provider,
            "anthropic_available": self.anthropic_client is not None,
            "openai_available": self.openai_client is not None,
            "supports_vision": True,
            "supports_embeddings": self.openai_client is not None,
            "prompt_manager_enabled": True,
            "available_analysis_types": [
                "unified",
                "modular",
                "metadata",
                "classification",
                "entities",
                "text",
                "design",
                "keywords",
                "communication",
            ],
        }

    def get_available_analysis_types(self) -> List[str]:
        """Get list of available analysis types"""
        return [
            "unified",  # Single comprehensive analysis
            "modular",  # Multiple specialized analyses
            "metadata",  # Core metadata only
            "classification",  # Classification only
            "entities",  # Entity extraction only
            "text",  # Text extraction only
            "design",  # Design elements only
            "keywords",  # Taxonomy keywords only
            "communication",  # Communication focus only
        ]

    async def analyze_text_chunk(
        self, text_chunk: str, filename: str, analysis_type: str = "unified"
    ) -> Dict[str, Any]:
        """
        Run AI analysis on a chunk of text.
        This is a lightweight version of the main analysis pipeline.
        """
        try:
            if analysis_type == "unified":
                prompt_data = await self.prompt_manager.get_unified_analysis_prompt(
                    filename
                )
                enhanced_prompt = self._enhance_prompt_with_text(
                    prompt_data["user"], text_chunk
                )
                analysis_result = await self._call_anthropic_api_with_system(
                    prompt_data["system"], enhanced_prompt
                )

                # Propagate API-level errors and unparseable responses so the caller
                # can set status=FAILED rather than silently storing bad data.
                if isinstance(analysis_result, dict) and "error" in analysis_result:
                    raise Exception(f"AI API call failed: {analysis_result['error']}")
                if isinstance(analysis_result, dict) and "raw_response" in analysis_result:
                    raise Exception("AI returned unparseable response (not valid JSON)")

                # Validate keyword mappings
                mappings = self._extract_mappings_from_analysis(analysis_result)
                validated_mappings = await self._validate_keyword_mappings(mappings)
                if "keyword_mappings" in analysis_result:
                    analysis_result["keyword_mappings"] = validated_mappings

                # Normalize the analysis result to ensure a consistent structure
                if "document_analysis" not in analysis_result:
                    analysis_result["document_analysis"] = {
                        "summary": analysis_result.get("summary", ""),
                        "document_type": analysis_result.get("document_type", ""),
                        "campaign_type": analysis_result.get("campaign_type", ""),
                        "document_tone": analysis_result.get("document_tone", ""),
                    }

                analysis_result["prompt_version"] = self.prompt_manager.PROMPT_VERSION
            else:
                # For simplicity, this example only implements the 'unified' chunk analysis
                logger.warning(
                    f"Analysis type '{analysis_type}' not fully supported for chunked processing. Using fallback."
                )
                analysis_result = {"summary": text_chunk[:100]}  # Basic fallback

            return analysis_result
        except Exception as e:
            logger.error(f"Error analyzing text chunk for {filename}: {str(e)}")
            raise

    def analyze_document_sync(
        self, file_path: str, filename: str, analysis_type: str = "unified"
    ) -> Dict[str, Any]:
        """Synchronous version of analyze_document"""
        return asyncio.run(self.analyze_document(file_path, filename, analysis_type))

    def generate_embeddings_sync(self, text: str) -> Optional[List[float]]:
        """Synchronous version of generate_embeddings"""
        return asyncio.run(self.generate_embeddings(text))

    def extract_text_from_pdf_sync_generator(
        self, file_content: bytes
    ) -> Generator[Tuple[int, str], None, None]:
        """Synchronous generator for extracting text from PDF pages."""

        async def get_all_pages():
            return [
                page
                async for page in self._extract_text_from_pdf_generator(file_content)
            ]

        try:
            pages = asyncio.run(get_all_pages())
            for page in pages:
                yield page
        except Exception as e:
            logger.error(f"Error in sync PDF generator: {e}")
            # Yield nothing if there's an error
            return

    def analyze_text_chunk_sync(
        self, text_chunk: str, filename: str, analysis_type: str = "unified"
    ) -> Dict[str, Any]:
        """Synchronous version of analyze_text_chunk."""
        return asyncio.run(self.analyze_text_chunk(text_chunk, filename, analysis_type))
