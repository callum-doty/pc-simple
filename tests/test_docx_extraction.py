"""
Tests for DOCX text extraction (services/ai_service.py) and the empty-content
safety guard in analyze_document().

Before this fix, uploading a .docx produced zero extracted text and no image
data (docx isn't in the image/pdf vision path), so the unified-analysis prompt
reached Claude with nothing to analyze but the filename. Claude would still
produce a plausible-looking JSON analysis — a fabricated summary and keyword
mappings grounded in nothing — which then got stored and embedded as if it
were real, with the document silently marked COMPLETED. These tests cover the
real python-docx extraction path and the guard that now fails such documents
loudly (ValueError -> document marked FAILED) instead of hallucinating.
"""

import io
from unittest.mock import AsyncMock, Mock

import pytest
from docx import Document as DocxDocument

from services.ai_service import AIService


def make_docx_bytes(paragraphs=None, table_rows=None) -> bytes:
    doc = DocxDocument()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row()
            for cell, value in zip(row.cells, row_values):
                cell.text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_ai_service() -> AIService:
    return AIService(db=Mock())


class TestExtractTextFromDocument:
    @pytest.mark.asyncio
    async def test_extracts_paragraph_text(self):
        content = make_docx_bytes(paragraphs=["First paragraph.", "Second paragraph."])
        ai_service = make_ai_service()

        text = await ai_service._extract_text_from_document(content)

        assert "First paragraph." in text
        assert "Second paragraph." in text

    @pytest.mark.asyncio
    async def test_extracts_table_text(self):
        content = make_docx_bytes(
            paragraphs=["Intro line."],
            table_rows=[["Client", "Amount"], ["Acme PAC", "$5,000"]],
        )
        ai_service = make_ai_service()

        text = await ai_service._extract_text_from_document(content)

        assert "Intro line." in text
        assert "Acme PAC" in text
        assert "$5,000" in text

    @pytest.mark.asyncio
    async def test_skips_empty_paragraphs(self):
        content = make_docx_bytes(paragraphs=["Real text.", "", "   "])
        ai_service = make_ai_service()

        text = await ai_service._extract_text_from_document(content)

        assert text == "Real text."

    @pytest.mark.asyncio
    async def test_returns_empty_string_for_corrupt_file(self):
        ai_service = make_ai_service()

        text = await ai_service._extract_text_from_document(b"not a real docx file")

        assert text == ""


class TestEmptyContentGuard:
    @pytest.mark.asyncio
    async def test_raises_for_document_type_with_no_extractable_text(self):
        ai_service = make_ai_service()
        ai_service.storage_service.get_file = AsyncMock(return_value=b"not a real docx")

        with pytest.raises(ValueError, match="No text could be extracted"):
            await ai_service.analyze_document("/fake/path.docx", "empty.docx")

    @pytest.mark.asyncio
    async def test_raises_for_empty_text_file(self):
        ai_service = make_ai_service()
        ai_service.storage_service.get_file = AsyncMock(return_value=b"   \n  ")

        with pytest.raises(ValueError, match="No text could be extracted"):
            await ai_service.analyze_document("/fake/path.txt", "empty.txt")

    @pytest.mark.asyncio
    async def test_does_not_raise_for_pdf_with_empty_text(self, monkeypatch):
        ai_service = make_ai_service()
        ai_service.storage_service.get_file = AsyncMock(return_value=b"%PDF-fake")
        monkeypatch.setattr(ai_service, "_extract_text", AsyncMock(return_value=""))
        monkeypatch.setattr(
            ai_service,
            "_perform_unified_analysis",
            AsyncMock(return_value={"summary": "ok"}),
        )

        result = await ai_service.analyze_document("/fake/path.pdf", "scan.pdf")

        assert result["ai_analysis"] == {"summary": "ok"}

    @pytest.mark.asyncio
    async def test_does_not_raise_for_image_with_empty_text(self, monkeypatch):
        ai_service = make_ai_service()
        ai_service.storage_service.get_file = AsyncMock(return_value=b"\x89PNGfake")
        monkeypatch.setattr(ai_service, "_extract_text", AsyncMock(return_value=""))
        monkeypatch.setattr(
            ai_service,
            "_perform_unified_analysis",
            AsyncMock(return_value={"summary": "ok"}),
        )

        result = await ai_service.analyze_document("/fake/path.png", "scan.png")

        assert result["ai_analysis"] == {"summary": "ok"}

    @pytest.mark.asyncio
    async def test_docx_with_real_text_does_not_raise(self, monkeypatch):
        content = make_docx_bytes(paragraphs=["Actual document content here."])
        ai_service = make_ai_service()
        ai_service.storage_service.get_file = AsyncMock(return_value=content)
        monkeypatch.setattr(
            ai_service,
            "_perform_unified_analysis",
            AsyncMock(return_value={"summary": "ok"}),
        )

        result = await ai_service.analyze_document("/fake/path.docx", "real.docx")

        assert result["extracted_text"] == "Actual document content here."
        assert result["ai_analysis"] == {"summary": "ok"}
